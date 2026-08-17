import sys
try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

doc_path = r"c:\Users\ASUS NUC\Desktop\Code\OME\HuongDanSuDung_OpenMediaSDK.docx"

try:
    doc = Document(doc_path)
except Exception as e:
    doc = Document()

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    from docx.shared import Pt
    r.font.size = Pt(16 if level==1 else 14)

def add_list_item(doc, text):
    p = doc.add_paragraph()
    p.add_run("• " + text)

add_heading(doc, 'HƯỚNG DẪN CHẠY THỬ SERVER VÀ DEMO', level=1)

add_heading(doc, '1. Chạy OpenMediaServer (Backend Engine)', level=2)
doc.add_paragraph('Đây là trái tim của hệ thống (kiến trúc Exhand), chạy độc lập và lắng nghe các lệnh từ Client (UI) thông qua IPC Named Pipes.')
add_list_item(doc, 'Mở PowerShell và di chuyển vào thư mục chứa file thực thi:')
doc.add_paragraph('  cd "C:\\Users\\ASUS NUC\\Desktop\\Code\\OME\\build-production\\bin\\Release"')
add_list_item(doc, 'Khởi động Server:')
doc.add_paragraph('  .\\OpenMediaServer.exe')
doc.add_paragraph('(Lúc này Server sẽ chạy ngầm, tạo một Named Pipe mặc định và chờ Client kết nối).')

add_heading(doc, '2. Chạy Pipeline Demo (Client App)', level=2)
doc.add_paragraph('Đây là một ứng dụng client độc lập mẫu. Ứng dụng này sẽ giả lập việc đọc một file video, đưa qua Mixer để xử lý và hiển thị lên một cửa sổ Preview đồng thời phát âm thanh ra loa.')
p = doc.add_paragraph()
p.add_run('Lưu ý quan trọng: ').bold = True
p.add_run('Mã nguồn của pipeline_demo mặc định sẽ tìm một file tên là ')
p.add_run('sample.mp4').italic = True
p.add_run(' trong thư mục hiện tại để phát.')

add_list_item(doc, 'Bạn cần copy hoặc tải một file video bất kỳ (định dạng mp4) và đổi tên thành sample.mp4, sau đó đặt vào thư mục build-production\\bin\\Release.')
add_list_item(doc, 'Khởi chạy demo:')
doc.add_paragraph('  .\\pipeline_demo.exe')
doc.add_paragraph('(Một cửa sổ Windows "OpenMedia Pipeline Preview" sẽ hiện lên, phát video và âm thanh từ file sample.mp4).')

add_heading(doc, '3. Các Demo Khác', level=2)
doc.add_paragraph('Trong thư mục còn có các bản demo khác:')
add_list_item(doc, 'rtmp_demo.exe: Đọc file và đẩy stream RTMP (ví dụ đẩy lên YouTube/Facebook).')
add_list_item(doc, 'ndi_srt_output.exe: Demo xuất tín hiệu video ra NDI hoặc truyền qua SRT.')
add_list_item(doc, 'simple_player.exe: Trình phát media cơ bản.')

p2 = doc.add_paragraph()
p2.add_run('Lưu ý: ').bold = True
p2.add_run('Để xem console log rõ ràng nhất khi chạy, hãy đảm bảo rằng file .env.production (hoặc .env.demo nếu bạn build demo) có bật OME_LOG_TO_CONSOLE=true.')

doc.save(doc_path)
print("Updated DOCX successfully!")
