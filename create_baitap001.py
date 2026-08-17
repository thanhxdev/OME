import sys
import os

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

out_dir = r"c:\Users\ASUS NUC\Desktop\Code\OME\HDThucHanh"
os.makedirs(out_dir, exist_ok=True)
doc_path = os.path.join(out_dir, "BaiTap001.docx")

doc = Document()

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16 if level==1 else 14)

def add_code(doc, code_text):
    p = doc.add_paragraph()
    r = p.add_run(code_text)
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
    # highlight with light grey background if possible, or just change text color
    r.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)
    
doc.add_heading('Bài Tập 001: Viết ứng dụng C++ đơn giản với OME SDK', 0)

add_heading(doc, 'Kiến trúc pipeline cơ bản của OME SDK', level=2)
doc.add_paragraph('1. Khởi tạo các thành phần (Components): Khởi tạo Source (ví dụ đọc file), Processor (ví dụ bộ giải mã Decoder), Sink (đầu ra như hiển thị ra màn hình hoặc ghi ra file).')
doc.add_paragraph('2. Khởi tạo & Cấu hình: Gọi Open() hoặc Initialize().')
doc.add_paragraph('3. Kết nối (Connect): Liên kết Source tới Decoder, Decoder tới Sink theo dạng chuỗi.')
doc.add_paragraph('4. Bắt đầu (Start): Khởi chạy các luồng xử lý bên trong của các component.')
doc.add_paragraph('5. Vòng lặp (Loop): Kéo (Pull) dữ liệu từ component cuối cùng trong pipeline để xử lý (hoặc để mặc cho pipeline tự chảy nếu là dạng push).')
doc.add_paragraph('6. Dọn dẹp (Cleanup): Dừng Stop() và Ngắt kết nối Disconnect().')

add_heading(doc, '1. File main.cpp', level=2)
code_cpp = """#include <iostream>
#include <memory>
#include <thread>
#include <chrono>

#include <openmedia/core/Logger.h>
#include <openmedia/io/FileSource.h>
#include <openmedia/codecs/FFmpegH264Decoder.h>

using namespace openmedia::core;
using namespace openmedia::io;
using namespace openmedia::codecs;

int main(int argc, char* argv[]) {
    Logger::SInfo("App", "Khởi động ứng dụng OME cơ bản...");

    if (argc < 2) {
        Logger::SError("App", "Cách dùng: hello_ome <duong_dan_file_video>");
        return 1;
    }

    std::string filePath = argv[1];

    auto fileSource = std::make_shared<FileSource>();
    auto decoder = std::make_shared<FFmpegH264Decoder>();

    auto resOpen = fileSource->Open(filePath);
    if (!resOpen.has_value()) {
        Logger::SError("App", "Không thể mở file video");
        return 1;
    }

    if (!fileSource->Initialize() || !decoder->Initialize()) {
        Logger::SError("App", "Khởi tạo thất bại.");
        return 1;
    }

    fileSource->Connect(decoder);

    decoder->Start();
    fileSource->Start();

    Logger::SInfo("App", "Pipeline đang chạy...");

    int frameCount = 0;
    while (fileSource->GetState() == PipelineState::Running && frameCount < 10) {
        std::this_thread::sleep_for(std::chrono::milliseconds(10));
        auto frameResult = decoder->PullFrame();
        if (frameResult.has_value() && frameResult.value()) {
            auto frame = frameResult.value();
            Logger::SInfo("App", "Frame {} - {}x{}", frameCount, frame->GetWidth(), frame->GetHeight());
            frameCount++;
        }
    }

    fileSource->Stop();
    decoder->Stop();
    fileSource->Disconnect();
    return 0;
}
"""
add_code(doc, code_cpp)

add_heading(doc, '2. File CMakeLists.txt', level=2)
code_cmake = """cmake_minimum_required(VERSION 3.28)

project(hello_ome LANGUAGES CXX)

add_executable(hello_ome main.cpp)

target_link_libraries(hello_ome PRIVATE
    OpenMedia.Core
    OpenMedia.IO
    OpenMedia.Codecs
)

set_target_properties(hello_ome PROPERTIES
    CXX_STANDARD 23
    CXX_STANDARD_REQUIRED ON
)
"""
add_code(doc, code_cmake)

doc.save(doc_path)
print(f"File saved successfully to {doc_path}")
