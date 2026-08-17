import sys
import os

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
from docx.shared import Pt

doc_path = r"c:\Users\ASUS NUC\Desktop\Code\OME\HDThucHanh\BaiTap001.docx"
doc = Document(doc_path)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16 if level==1 else 14)

add_heading(doc, '3. Hướng dẫn Tích hợp, Biên dịch và Chạy thử', level=2)

doc.add_paragraph('Sau khi bạn đã viết xong 2 file main.cpp và CMakeLists.txt (đặt trong cùng một thư mục tên là hello_ome), hãy làm theo các bước sau để tích hợp vào project OME và biên dịch:')
doc.add_paragraph('Bước 1: Copy toàn bộ thư mục hello_ome của bạn vào trong: c:\\Users\\ASUS NUC\\Desktop\\Code\\OME\\samples\\cpp\\')
doc.add_paragraph('Bước 2: Mở file c:\\Users\\ASUS NUC\\Desktop\\Code\\OME\\samples\\cpp\\CMakeLists.txt và thêm dòng chữ sau xuống vị trí dưới cùng của file:')
doc.add_paragraph('  add_subdirectory(hello_ome)')
doc.add_paragraph('Bước 3: Mở PowerShell, di chuyển vào thư mục gốc của dự án OME và chạy lệnh build SDK (lệnh này sẽ tự động biên dịch lại thư viện và build luôn cả code ví dụ hello_ome của bạn):')
doc.add_paragraph('  cd "C:\\Users\\ASUS NUC\\Desktop\\Code\\OME"')
doc.add_paragraph('  .\\tools\\scripts\\build.ps1 -Environment production')
doc.add_paragraph('Bước 4: Chạy thử ứng dụng')
doc.add_paragraph('Sau khi build thành công, file thực thi sẽ được đặt trong thư mục Release. Bạn hãy chuẩn bị một file video (ví dụ test.mp4) và chạy lệnh sau để xem kết quả:')
doc.add_paragraph('  .\\build-production\\bin\\Release\\hello_ome.exe test.mp4')

doc.save(doc_path)
print(f"File updated successfully at {doc_path}")
